import os
import numpy as np
import nltk
import pandas as pd
from sklearn.decomposition import PCA
from sklearn.cluster import KMeans
import matplotlib.pyplot as plt
from docx import Document
import warnings
from gensim.models import Word2Vec
from win32com.client import Dispatch
from sklearn.metrics import silhouette_score, davies_bouldin_score

# Ignore all warnings
warnings.filterwarnings("ignore")


# Function to preprocess the text
def preprocess(text):
    tokens = nltk.word_tokenize(text.lower())
    stopwords = set(nltk.corpus.stopwords.words('english'))
    tokens = [token for token in tokens if token.isalpha() and token not in stopwords]
    return tokens


# Function to create the co-occurrence matrix
def create_co_occurrence_matrix(documents, window_size=2):
    vocab = set()
    processed_docs = []

    for doc in documents:
        tokens = preprocess(doc)
        processed_docs.append(tokens)
        vocab.update(tokens)

    vocab = list(vocab)
    vocab_size = len(vocab)
    word_to_id = {word: i for i, word in enumerate(vocab)}

    # Initialize the co-occurrence matrix
    co_occurrence_matrix = np.zeros((vocab_size, vocab_size), dtype=np.int32)

    # Populate the co-occurrence matrix
    for tokens in processed_docs:
        for i, token in enumerate(tokens):
            token_id = word_to_id[token]
            start = max(0, i - window_size)
            end = min(len(tokens), i + window_size + 1)
            for j in range(start, end):
                if i != j:
                    co_occurrence_token_id = word_to_id[tokens[j]]
                    co_occurrence_matrix[token_id, co_occurrence_token_id] += 1

    return co_occurrence_matrix, vocab


# Function to read documents from .docx files
def read_documents(file_paths):
    documents = []
    for file_path in file_paths:
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            documents.append('\n'.join(full_text))
        elif file_path.endswith('.doc'):
            # Convert .doc to .docx using pywin32
            word = Dispatch('Word.Application')
            doc = word.Documents.Open(file_path)
            docx_path = file_path + 'x'
            doc.SaveAs2(docx_path, FileFormat=16)  # Save as .docx format
            doc.Close()
            word.Quit()
            docx = Document(docx_path)
            full_text = []
            for para in docx.paragraphs:
                full_text.append(para.text)
            for table in docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            documents.append('\n'.join(full_text))
            os.remove(docx_path)  # Clean up the temporary .docx file
        elif file_path.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfFileReader(open(file_path, 'rb'))
            full_text = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                full_text.append(page.extract_text())
            documents.append('\n'.join(full_text))
    return documents



# Collect folder path from user input
folder_path = input("Enter the path to the folder containing .docx files: ").strip('"')

# Get all .docx files in the specified folder
file_paths = [os.path.join(folder_path, file) for file in os.listdir(folder_path) if file.endswith('.docx')]

# Read documents from files
documents = read_documents(file_paths)

# Lists to store results for each document
all_words_list = []
processed_docs = []

# Process each document individually
for doc_index, document in enumerate(documents):
    print(f"\nProcessing Document {doc_index + 1}...")

    # Create the co-occurrence matrix for the document
    co_occurrence_matrix, vocab = create_co_occurrence_matrix([document])

    # Display the co-occurrence matrix
    co_occurrence_df = pd.DataFrame(co_occurrence_matrix, index=vocab, columns=vocab)
    print("\nCo-occurrence Matrix:")
    print(co_occurrence_df.to_string())

    # Identify pairs of words with the highest co-occurrence counts
    pairs = []
    for i in range(len(vocab)):
        for j in range(i + 1, len(vocab)):
            if co_occurrence_matrix[i, j] > 0:
                pairs.append((vocab[i], vocab[j], co_occurrence_matrix[i, j]))

    # Sort pairs by co-occurrence count in descending order
    pairs.sort(key=lambda x: x[2], reverse=True)

    # Find the minimum co-occurrence count
    if pairs:
        min_count = min(pairs, key=lambda x: x[2])[2]

    # Identify and remove pairs with the minimum co-occurrence count
    min_pairs = [pair for pair in pairs if pair[2] == min_count]
    pairs = [pair for pair in pairs if pair[2] > min_count]

    # Add words from the minimum co-occurrence pairs to singleton words list
    singleton_words = set()
    for pair in min_pairs:
        singleton_words.update(pair[:2])

    # Collect words that are part of higher co-occurrence pairs
    high_co_occurrence_words = set()
    for pair in pairs:
        high_co_occurrence_words.update(pair[:2])

    # Add words with no co-occurrences to the singleton words list
    for i, word in enumerate(vocab):
        if not np.any(co_occurrence_matrix[i]):
            singleton_words.add(word)

    # Remove words that are part of high co-occurrence pairs from singleton words list
    singleton_words = singleton_words - high_co_occurrence_words

    # Display the pairs of words with the highest co-occurrence counts
    print("\nPairs of Words with the Highest Co-occurrence Counts:")
    for pair in pairs:
        print(f"{pair[0]} - {pair[1]}: {pair[2]}")

    # Display singleton words
    print("\nSingleton Words (based on co-occurrence):")
    for word in singleton_words:
        print(word)

    # Vectorizing singleton words using Word2Vec
    singleton_docs = [list(singleton_words)]

    # Create Word2Vec model
    w2v_model = Word2Vec(singleton_docs, vector_size=100, window=2, min_count=1, sg=1)

    # Create vectors for singleton words
    word_vectors = np.array([w2v_model.wv[word] for word in singleton_words])

    # PCA on Word2Vec
    pca_w2v = PCA(n_components=min(word_vectors.shape[0], word_vectors.shape[1]))
    X_w2v_pca = pca_w2v.fit_transform(word_vectors)

    # Get the explained variance ratio
    explained_variance_ratio_w2v = pca_w2v.explained_variance_ratio_

    # Find the component with the least explained variance
    least_variance_component_w2v = np.argmin(explained_variance_ratio_w2v)

    # Find the corresponding word with the least variance
    least_occuring_word_w2v = list(singleton_words)[least_variance_component_w2v]

    print(f"Least occurring word in Word2Vec PCA: {least_occuring_word_w2v}")

    # Collect all words for the document
    all_words = list(high_co_occurrence_words) + [least_occuring_word_w2v]
    all_words_list.append(all_words)
    print(f"All words for Document {doc_index + 1}: {all_words}")

    # Store the processed document as a single string
    processed_docs.append(" ".join(all_words))

# Vectorize all documents using Word2Vec
all_docs = [doc.split() for doc in processed_docs]

# Create Word2Vec model
w2v_model = Word2Vec(all_docs, vector_size=100, window=2, min_count=1, sg=1)

# Create vectors for all documents
doc_vectors = np.array([np.mean([w2v_model.wv[word] for word in doc], axis=0) for doc in all_docs])

# Elbow method to determine the optimal number of clusters
sum_of_squared_distances = []
K = range(1, min(len(file_paths), 10))  # Adjust the range to avoid ValueError
for k in K:
    km = KMeans(n_clusters=k)
    km = km.fit(doc_vectors)
    sum_of_squared_distances.append(km.inertia_)

plt.figure(figsize=(12, 6))
plt.plot(K, sum_of_squared_distances, 'bx-')
plt.xlabel('Number of clusters')
plt.ylabel('Sum of squared distances')
plt.title('Elbow Method For Optimal k')
plt.show()

# User inputs the desired number of clusters based on the elbow graph
while True:
    n_clusters = int(input(f"Enter the number of clusters (1 to {min(len(file_paths), 9)}): "))
    if 1 <= n_clusters <= min(len(file_paths), 9):
        break
    else:
        print(f"Please enter a number between 1 and {min(len(file_paths), 9)}.")

# Perform KMeans clustering with the user-defined number of clusters
kmeans = KMeans(n_clusters=n_clusters)
kmeans.fit(doc_vectors)
labels = kmeans.labels_

# Perform PCA to reduce dimensions for plotting
pca = PCA(n_components=2)
doc_vectors_pca = pca.fit_transform(doc_vectors)

# Plotting KMeans clusters on the same graph
plt.figure(figsize=(12, 6))
plt.scatter(doc_vectors_pca[:, 0], doc_vectors_pca[:, 1], c=labels, cmap='viridis')
plt.title('KMeans Clustering of All Documents')
plt.xlabel('PCA Component 1')
plt.ylabel('PCA Component 2')
plt.show()

for file_path, label in zip(file_paths, labels):
    document_name = os.path.splitext(os.path.basename(file_path))[0]  # Extracts 'Doc1' from 'F:\Doc1.docx'
    print(f"Document '{document_name}' is in Cluster {label}")

# Calculate evaluation metrics
silhouette_avg = silhouette_score(doc_vectors, labels)
davies_bouldin = davies_bouldin_score(doc_vectors, labels)
print(f"Silhouette Score: {silhouette_avg}")
print(f"Davies-Bouldin Index: {davies_bouldin}")